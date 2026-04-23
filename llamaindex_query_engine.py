"""
llamaindex_query_engine.py — LlamaIndex-based vector store query engine.

Wraps LlamaIndex VectorStoreIndex with Azure OpenAI LLM + embeddings to provide
a simple add-documents / query interface consistent with this repo's patterns.

QUICKSTART
----------
    from llamaindex_query_engine import LlamaQueryEngine

    engine = LlamaQueryEngine(system_prompt="You are an AmEx expert.")
    engine.add_file("docs/product.pdf")
    engine.build()

    response = engine.query("What is the annual fee?")
    print(response.response)

DEPENDENCIES
------------
    pip install llama-index llama-index-llms-azure-openai llama-index-embeddings-azure-openai pypdf python-docx
"""
import os
from pathlib import Path
from typing import Optional

from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.core.query_engine import BaseQueryEngine
from llama_index.llms.azure_openai import AzureOpenAI
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding


__all__ = ["LlamaQueryEngine"]


DEFAULT_SYSTEM = (
    "Answer using ONLY the context below. If the answer isn't in the context, "
    "say you don't know. Cite sources inline as [Source: <filename>]."
)

# Azure deployment names — override via constructor kwargs or env vars.
_DEFAULT_LLM_DEPLOYMENT = os.getenv("AZURE_LLM_DEPLOYMENT", "gpt-41")
_DEFAULT_EMBED_DEPLOYMENT = os.getenv("AZURE_EMBED_DEPLOYMENT", "text-embedding-3-large")
_DEFAULT_AZURE_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "")
_DEFAULT_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2024-06-01")


def _parse_file(path: str) -> str:
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    ext = Path(path).suffix.lower()
    if ext in {".txt", ".md"}:
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        from pypdf import PdfReader
        return "\n".join(p.extract_text() or "" for p in PdfReader(path).pages)
    if ext == ".docx":
        from docx import Document as DocxDocument
        return "\n".join(p.text for p in DocxDocument(path).paragraphs)
    raise ValueError(f"Unsupported file type: {ext}. Supported: .txt, .md, .pdf, .docx")


class LlamaQueryEngine:
    """Vector store query engine backed by LlamaIndex + Azure OpenAI.

    Usage:
        engine = LlamaQueryEngine()
        engine.add_file("product.pdf")
        engine.build()
        result = engine.query("What is the annual fee?")
    """

    def __init__(
        self,
        system_prompt: str = DEFAULT_SYSTEM,
        azure_endpoint: str = _DEFAULT_AZURE_ENDPOINT,
        api_key: Optional[str] = None,
        api_version: str = _DEFAULT_API_VERSION,
        llm_deployment: str = _DEFAULT_LLM_DEPLOYMENT,
        embed_deployment: str = _DEFAULT_EMBED_DEPLOYMENT,
        similarity_top_k: int = 4,
    ) -> None:
        self._system_prompt = system_prompt
        self._similarity_top_k = similarity_top_k
        self._documents: list[Document] = []
        self._engine: Optional[BaseQueryEngine] = None

        resolved_key = api_key or os.getenv("AZURE_OPENAI_API_KEY", "")

        Settings.llm = AzureOpenAI(
            model="gpt-4",
            deployment_name=llm_deployment,
            azure_endpoint=azure_endpoint,
            api_key=resolved_key,
            api_version=api_version,
            system_prompt=system_prompt,
        )

        Settings.embed_model = AzureOpenAIEmbedding(
            model="text-embedding-3-large",
            deployment_name=embed_deployment,
            azure_endpoint=azure_endpoint,
            api_key=resolved_key,
            api_version=api_version,
        )

    # ------------------------------------------------------------------
    # Document ingestion
    # ------------------------------------------------------------------

    def add_document(self, text: str, metadata: Optional[dict] = None) -> None:
        """Add a raw text document with optional metadata."""
        self._documents.append(Document(text=text, metadata=metadata or {}))
        self._engine = None  # invalidate built engine

    def add_file(self, path: str, extra_metadata: Optional[dict] = None) -> None:
        """Parse a local file and add it as a document."""
        text = _parse_file(path)
        metadata = {"source": Path(path).name, **(extra_metadata or {})}
        self.add_document(text, metadata)

    # ------------------------------------------------------------------
    # Index / engine lifecycle
    # ------------------------------------------------------------------

    def build(self) -> None:
        """Build (or rebuild) the vector index from all added documents."""
        if not self._documents:
            raise ValueError("No documents added. Call add_document() or add_file() first.")
        index = VectorStoreIndex.from_documents(self._documents)
        self._engine = index.as_query_engine(similarity_top_k=self._similarity_top_k)

    # ------------------------------------------------------------------
    # Querying
    # ------------------------------------------------------------------

    def query(self, question: str):
        """Run a query against the vector store. Returns a LlamaIndex Response object."""
        if self._engine is None:
            raise RuntimeError("Engine not built. Call build() after adding documents.")
        return self._engine.query(question)
